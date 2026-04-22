"""
Contract Filler для ИИ-юриста СЗ Дело
Автоматическое заполнение шаблонов договоров
"""
import os
import logging
from typing import Dict, Optional
from pathlib import Path
from datetime import datetime

from docxtpl import DocxTemplate
from docx import Document

logger = logging.getLogger(__name__)


class ContractFiller:
    """Заполнение шаблонов договоров"""
    
    def __init__(self, templates_dir: str = "data/templates"):
        self.templates_dir = Path(templates_dir)
        self._ensure_templates_dir()
    
    def _ensure_templates_dir(self):
        """Создание директории шаблонов"""
        self.templates_dir.mkdir(parents=True, exist_ok=True)
        
        # Создание примеров шаблонов если пусто
        if not any(self.templates_dir.glob("*.docx")):
            self._create_sample_templates()
    
    def _create_sample_templates(self):
        """Создание примеров шаблонов"""
        # Пример: договор подряда
        sample_contract = {
            'filename': 'construction_contract_template.docx',
            'content': {
                'contract_number': '№ {{ contract_number }}',
                'contract_date': 'от {{ contract_date }}',
                'customer_name': '{{ customer_name }}',
                'customer_inn': '{{ customer_inn }}',
                'contractor_name': '{{ contractor_name }}',
                'contractor_inn': '{{ contractor_inn }}',
                'object_address': '{{ object_address }}',
                'work_description': '{{ work_description }}',
                'contract_price': '{{ contract_price }}',
                'start_date': '{{ start_date }}',
                'end_date': '{{ end_date }}',
            }
        }
        
        logger.info(f"Sample templates directory created: {self.templates_dir}")
    
    def get_available_templates(self) -> list:
        """Получение списка доступных шаблонов"""
        templates = []
        
        for file in self.templates_dir.glob("*.docx"):
            templates.append({
                'filename': file.name,
                'path': str(file),
                'name': file.stem.replace('_', ' ').title()
            })
        
        return templates
    
    def fill_template(self, 
                     template_path: str, 
                     data: Dict,
                     output_path: str = None) -> str:
        """
        Заполнение шаблона данными
        
        Args:
            template_path: Путь к шаблону DOCX
            data: Данные для заполнения (ключи должны совпадать с переменными в шаблоне)
            output_path: Путь для сохранения результата
            
        Returns:
            Путь к заполненному файлу
        """
        try:
            # Загрузка шаблона
            tpl = DocxTemplate(template_path)
            
            # Добавление стандартных полей
            if 'contract_date' not in data:
                data['contract_date'] = datetime.now().strftime('%d.%m.%Y')
            
            if 'current_year' not in data:
                data['current_year'] = datetime.now().year
            
            # Рендеринг
            tpl.render(data)
            
            # Сохранение
            if not output_path:
                output_path = str(self.templates_dir / f"filled_{Path(template_path).name}")
            
            tpl.save(output_path)
            
            logger.info(f"Template filled: {output_path}")
            return output_path
        
        except Exception as e:
            logger.error(f"Template fill error: {e}")
            raise
    
    def fill_and_analyze(self, 
                        template_path: str, 
                        data: Dict,
                        analyze: bool = True) -> tuple:
        """
        Заполнение шаблона с последующим анализом
        
        Args:
            template_path: Путь к шаблону
            data: Данные для заполнения
            analyze:是否需要分析
            
        Returns:
            (путь к файлу, анализ рисков)
        """
        from backend.risk_analyzer import create_risk_analyzer
        from backend.document_parser import create_parser
        
        # Заполнение
        output_path = self.fill_template(template_path, data)
        
        if not analyze:
            return output_path, None
        
        # Анализ
        parser = create_parser()
        analyzer = create_risk_analyzer()
        
        # Чтение заполненного файла
        doc = Document(output_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        
        # Анализ рисков
        risks = analyzer.analyze_contract(text)
        
        return output_path, risks
    
    def create_template_from_contract(self, 
                                     contract_path: str,
                                     template_name: str) -> str:
        """
        Создание шаблона из существующего договора
        
        Args:
            contract_path: Путь к договору
            template_name: Имя шаблона
            
        Returns:
            Путь к шаблону
        """
        import re
        
        parser = create_parser()
        text = parser.parse_file(contract_path)
        
        # Замена конкретных значений на переменные
        replacements = {
            r'\d{1,2}\.\d{1,2}\.\d{4}': '{{ date }}',
            r'№\s*\d+': '{{ contract_number }}',
            r'\d{1,3}\s*(?:млн|тыс)\.\s*\d+': '{{ amount }}',
            r'г\.\s*Москва': '{{ city }}',
        }
        
        template_text = text
        for pattern, replacement in replacements.items():
            template_text = re.sub(pattern, replacement, template_text, flags=re.IGNORECASE)
        
        # Сохранение шаблона
        template_path = self.templates_dir / f"{template_name}.docx"
        
        # Создание DOCX
        doc = Document()
        for paragraph in template_text.split('\n\n'):
            doc.add_paragraph(paragraph)
        
        doc.save(str(template_path))
        
        logger.info(f"Template created: {template_path}")
        return str(template_path)


# ============================================================================
# Фабрика
# ============================================================================

def create_filler() -> ContractFiller:
    """Создание заполнителя договоров"""
    return ContractFiller()
